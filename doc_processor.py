import os
import re
from typing import List
import pdfplumber
import docx
from langchain.schema import Document as LangchainDocument
from langchain.prompts import ChatPromptTemplate
from langchain_community.chat_models import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough

class DocumentProcessor:
    def __init__(self, openai_api_key: str):
        # Initialize the LLM with a higher temperature for more creative restructuring
        self.llm = ChatOpenAI(
            model="gpt-4-turbo",
            openai_api_key=openai_api_key,
            temperature=0.2 # Slightly higher temp for creativity in structuring
        )
        # Initialize a text splitter to handle long documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=8000,  # Adjust based on model context window
            chunk_overlap=200,
            length_function=len,
        )

    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extracts text from various file types."""
        if file_type == "pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            return self.extract_text_from_docx(file_path)
        elif file_type == "txt":
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extracts text from a PDF file."""
        print(f"Extracting text from PDF: {file_path}")
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text from each page
                    page_text = page.extract_text()
                    if page_text:
                        # Clean up text: replace multiple newlines/whitespaces
                        cleaned_text = re.sub(r'\s+', ' ', page_text.strip())
                        text += cleaned_text + " "
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
        return text.strip()

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extracts text from a DOCX file."""
        print(f"Extracting text from DOCX: {file_path}")
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extracts text from a TXT file."""
        print(f"Extracting text from TXT: {file_path}")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

    def structure_document(self, text: str) -> str:
        """
        The core function that uses the LLM to structure the document.
        This implements a Map-Reduce strategy for long documents.
        """
        # Define our sophisticated prompt template
        prompt_template = """
        You are an expert editor and technical writer. Your task is to transform unstructured text into a perfectly organized, easy-to-read document.

        Follow these guidelines strictly:
        1. Analyze the input text to identify main topics, sections, and key points.
        2. Create a logical hierarchy using Markdown formatting:
           - Use # Headers for main titles
           - Use ## Subheaders for major sections
           - Use ### Subsubheaders for sub-sections
           - Use bullet points (- or *) for lists and key points
           - Use **bold** for key terms and important concepts
           - Use italics for definitions or subtle emphasis
        3. Maintain all crucial information from the source text.
        4. Improve readability by breaking down long paragraphs and removing redundancy.
        5. Ensure the output is comprehensive yet concise.

        Here is the text to structure:
        {text}
        """

        # If the document is short, process it in one go
        if len(text) < 6000:
            prompt = ChatPromptTemplate.from_template(prompt_template)
            chain = prompt | self.llm | StrOutputParser()
            return chain.invoke({"text": text})
        
        # For long documents: Implement Map-Reduce
        else:
            print("Document is long. Using Map-Reduce strategy...")
            # Split the document into manageable chunks
            texts = self.text_splitter.split_text(text)
            docs = [LangchainDocument(page_content=t) for t in texts]
            
            # Map step: Structure each chunk individually
            map_prompt = ChatPromptTemplate.from_template(
                "Extract and structure the key information from this section of a document. "
                "Use clear headings and bullet points. Here is the section:\n\n{text}"
            )
            map_chain = map_prompt | self.llm | StrOutputParser()
            
            # Reduce step: Combine all structured chunks coherently
            reduce_prompt = ChatPromptTemplate.from_template(
                "You are synthesizing a complete structured document from multiple sections. "
                "Combine these structured sections into a single, coherent, well-organized document. "
                "Ensure consistent formatting and logical flow throughout. "
                "Sections to combine:\n\n{text}"
            )
            reduce_chain = reduce_prompt | self.llm | StrOutputParser()
            
            # Run the map-reduce process
            structured_chunks = map_chain.batch([{"text": doc.page_content} for doc in docs])
            combined_text = "\n\n".join(structured_chunks)
            
            # Final reduction to create a unified document
            return reduce_chain.invoke({"text": combined_text})